from pathlib import Path

from docx import Document

from . import docx_fun

BASE_DIR = Path(__file__).resolve().parent.parent
NOTICE_TEMPLATE = BASE_DIR / "templates" / "模板-漏洞应急响应通告.docx"


def solution(image_filename, uploadCheck, get_version):
    """Prepare the main notice template and inject the optional version screenshot."""
    doc = Document(str(NOTICE_TEMPLATE))
    if uploadCheck == "true":
        version1 = "1、如何检测组件系统版本"
        version3 = "2、官方修复建议"
        version4 = "3、临时修复建议"
        index = docx_fun.find_paragraph_with_string(doc, "version2")
        target_paragraph = doc.paragraphs[index + 1]
        if image_filename:
            target_paragraph.add_run().add_picture(image_filename)
    else:
        docx_fun.delete_paragraph_index(doc, "version1", 3)
        version1 = ""
        version3 = "1、官方修复建议"
        version4 = "2、临时修复建议"

    context = {
        "version1": version1,
        "version2": get_version,
        "version3": version3,
        "version4": version4,
    }
    return context, doc
